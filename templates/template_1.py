import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_cv_from_json(json_file, output_docx):
    # Load JSON data
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create a new Word document
    doc = Document()

    # Set custom margins (slightly more padding on the left and right)
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(15)   # Slight padding on the left
        section.right_margin = Pt(15)  # Slight padding on the right
        section.top_margin = Pt(10)    # Small top margin
        section.bottom_margin = Pt(10) # Small bottom margin

    # Set global styling
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)  # Use a smaller font size for compactness

    # Add name as the title
    title = doc.add_heading(level=1)
    run = title.add_run(data.get("name", "Your Name"))
    run.bold = True
    run.font.size = Pt(14)  # Smaller title size
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(0)  # Reduce space after title

    # Add contact information
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_info.add_run(data.get("contact", "Your Contact Information") + " | ").bold = True
    contact_info.add_run(data.get("email", "Your Email") + " | ").bold = True
    contact_info.add_run(data.get("phone", "Your Phone Number")).bold = True
    contact_info.paragraph_format.space_after = Pt(0)  # Remove space after contact

    # Add Summary
    if "summary" in data:
        doc.add_heading("Summary", level=2)
        summary_paragraph = doc.add_paragraph(data["summary"])
        summary_paragraph.paragraph_format.space_after = Pt(0)  # Reduce space after summary

    # Add Work Experience
    if "experience" in data:
        doc.add_heading("Work Experience", level=2)
        for job in data["experience"]:
            job_title = doc.add_paragraph()
            run = job_title.add_run(job["title"] + " - " + job["company"])
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(job['dates']).italic = True
            doc.add_paragraph(job["description"])
            doc.add_paragraph().paragraph_format.space_after = Pt(0)  # Reduce space after each experience

    # Add Education
    if "education" in data:
        doc.add_heading("Education", level=2)
        for edu in data["education"]:
            edu_title = doc.add_paragraph()
            run = edu_title.add_run(edu["degree"] + " - " + edu["institution"])
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(edu['year']).italic = True
            doc.add_paragraph().paragraph_format.space_after = Pt(0)  # Reduce space after each education entry

    # Add Skills
    if "skills" in data:
        doc.add_heading("Skills", level=2)
        skills_paragraph = doc.add_paragraph()
        for skill in data["skills"]:
            run = skills_paragraph.add_run("• " + skill + "  ")
            run.bold = True
        skills_paragraph.paragraph_format.space_after = Pt(0)  # Reduce space after skills section

    # Save the document
    doc.save(output_docx)
    print(f"CV saved as {output_docx}")

if __name__ == "__main__":
    input_json = os.getenv("INPUT_JSON", "cv_data.json")
    output_docx = os.getenv("OUTPUT_DOCX", "output_cv.docx")
    create_cv_from_json(input_json, output_docx)
