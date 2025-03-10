import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
import os

def add_horizontal_line(doc):
    """Adds a horizontal line to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("\u2014" * 50)  # Creates a long horizontal line
    run.bold = True

def create_cv_from_json(json_file, output_docx):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    
    # Name at the top
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = name.add_run(data.get("name", "Your Name").upper())
    run.bold = True
    run.font.size = Pt(14)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run(data.get("phone", "Your Phone Number") + " • ").bold = True
    contact.add_run(data.get("email", "Your Email") + " • ").bold = True
    contact.add_run(data.get("location", "Your Location")).bold = True
    
    add_horizontal_line(doc)
    
    # Summary
    if "summary" in data:
        summary_title = doc.add_paragraph("SOFTWARE DEVELOPER")
        summary_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        summary_title.runs[0].bold = True
        
        summary = doc.add_paragraph(data["summary"])
        summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    add_horizontal_line(doc)
    
    # Strengths and Expertise
    if "skills" in data:
        skills_title = doc.add_paragraph("STRENGTHS AND EXPERTISE")
        skills_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        skills_title.runs[0].bold = True
        
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row = table.rows[0].cells
        
        for i, skill in enumerate(data["skills"][:9]):  # Limit to 9 for spacing
            row[i % 3].text += skill + "\n"
        
    add_horizontal_line(doc)
    
    # Experience Section
    if "experience" in data:
        exp_title = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        exp_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        exp_title.runs[0].bold = True
        
        for job in data["experience"]:
            job_title = doc.add_paragraph()
            run = job_title.add_run(job["title"] + "\n" + job["company"])
            run.bold = True
            job_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            dates = doc.add_paragraph(job["dates"])
            dates.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            dates.runs[0].italic = True
            
            doc.add_paragraph(job["description"])
            
            for acc in job.get("accomplishments", []):
                bullet = doc.add_paragraph("• " + acc)
                bullet.paragraph_format.space_after = Pt(0)
    
    doc.save(output_docx)
    print(f"CV saved as {output_docx}")

if __name__ == "__main__":
    input_json = os.getenv("INPUT_JSON", "cv_data.json")
    output_docx = os.getenv("OUTPUT_DOCX", "output_cv.docx")
    create_cv_from_json(input_json, output_docx)
