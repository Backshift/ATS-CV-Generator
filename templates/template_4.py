import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_horizontal_line(doc):
    """Adds a single horizontal line to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("\u2014" * 50)  # Creates a long horizontal line
    run.bold = True

def add_section_title(doc, title_text):
    """Adds a centered, bold section title with a horizontal line."""
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(title_text.upper())
    run.bold = True
    add_horizontal_line(doc)

def add_text_section(doc, title, content, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Adds a text section with a title and content."""
    if content:
        sec_title = doc.add_paragraph(title)
        sec_title.alignment = alignment
        sec_title.runs[0].bold = True
        
        sec_content = doc.add_paragraph(content)
        sec_content.alignment = alignment

def add_skills_section(doc, skills):
    """Adds a skills section with three columns."""
    if skills:
        add_section_title(doc, "STRENGTHS AND EXPERTISE")
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row = table.rows[0].cells
        
        for i, skill in enumerate(skills[:9]):  # Limit to 9 for spacing
            row[i % 3].text += skill + "\n"

def add_experience_section(doc, experience):
    """Adds a professional experience section."""
    if experience:
        add_section_title(doc, "PROFESSIONAL EXPERIENCE")
        
        for job in experience:
            job_title = doc.add_paragraph()
            job_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            run = job_title.add_run(job["title"] + " - " + job["company"])
            run.bold = True
            
            date_run = job_title.add_run(" " * 50 + job["dates"])
            date_run.bold = True
            
            doc.add_paragraph(job["description"])
            
            for acc in job.get("accomplishments", []):
                bullet = doc.add_paragraph("• " + acc)
                bullet.paragraph_format.space_after = Pt(0)

def create_cv_from_json(json_file, output_docx):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    
    # Name at the top
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = name.add_run(data.get("name", "Your Name").upper())
    run.bold = True
    run.font.size = Pt(14)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run(data.get("phone", "Your Phone Number") + " • ").bold = True
    contact.add_run(data.get("email", "Your Email") + " • ").bold = True
    contact.add_run(data.get("location", "Your Location")).bold = True
    
    add_horizontal_line(doc)
    
    # Summary
    add_text_section(doc, "SOFTWARE DEVELOPER", data.get("summary", ""), WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    add_skills_section(doc, data.get("skills", []))
    
    add_experience_section(doc, data.get("experience", []))
    
    doc.save(output_docx)
    print(f"CV saved as {output_docx}")

if __name__ == "__main__":
    input_json = os.getenv("INPUT_JSON", "cv_data.json")
    output_docx = os.getenv("OUTPUT_DOCX", "output_cv.docx")
    create_cv_from_json(input_json, output_docx)