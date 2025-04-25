import os
import docx
import datetime
import re
from docx.shared import RGBColor
from spellchecker import SpellChecker
from nltk.tokenize import word_tokenize
import nltk

from docx.shared import Pt

nltk.download('punkt')
nltk.download('punkt_tab')

# Function to read Word documents
def read_docx(file_path):
    doc = docx.Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

# Check file type (Word/RTF)
def check_file_type(file_path):
    return "✅ File type is acceptable." if file_path.lower().endswith(('.docx', '.rtf')) else "❌ Please use a Word or RTF document."

# Check last edited date
def check_last_modified_date(file_path):
    last_modified = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
    return "✅ Your document was last updated recently (within 2 months)." if (datetime.datetime.today() - last_modified).days < 60 else "❌ Your document was last updated more than 2 months ago. Consider updating."

# Check file name format
def check_file_name(file_path, full_name):
    file_name = os.path.basename(file_path)
    name_check = "✅ Good job, your name is in the file name!" if full_name.lower() in file_name.lower() else "❌ Consider adding your name to the file name."
    name_length_check = "✅ File name length is good." if len(file_name) <= 24 else "❌ Your file name is too long. Keep it concise."
    return name_check, name_length_check

# Check file size
def check_file_size(file_path):
    return "✅ File size is acceptable." if os.path.getsize(file_path) < 1048576 else "❌ Your file size is too large (over 1MB)."

# Count words and estimate pages
def count_words_and_pages(text):
    words = [word for word in word_tokenize(text) if word.isalpha()]
    word_count = len(words)
    page_count = max(1, word_count // 400)  # Approx. 400 words per page
    return word_count, page_count

# Check font consistency
def check_font_and_colors(file_path):
    doc = docx.Document(file_path)
    font_sizes, font_colors = set(), set()
    standard_fonts = ['Arial', 'Calibri', 'Times New Roman']
    font_ok = True
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font:
                if run.font.size:
                    font_sizes.add(run.font.size.pt)
                if run.font.color and run.font.color.rgb:
                    font_colors.add(run.font.color.rgb)
                if run.font.name and run.font.name not in standard_fonts:
                    font_ok = False
    
    color_check = "❌ Multiple font colors detected. Use only black." if len(font_colors) > 1 else "✅ Font color is consistent."
    size_check = "✅ Font sizes are consistent." if len(font_sizes) <= 2 else "❌ Multiple font sizes detected. Keep it consistent."
    font_check = "✅ Standard fonts used." if font_ok else "❌ Non-standard fonts detected. Use Arial, Calibri, or Times New Roman."
    
    return size_check, font_check, color_check

# Check required sections
def check_sections(text):
    required_sections = ['experience', 'education', 'skills']
    missing_sections = [section for section in required_sections if section not in text.lower()]
    return "✅ All key sections included." if not missing_sections else f"❌ Missing sections: {', '.join(missing_sections)}."

# Check contact information
def check_contact_info(text):
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    phone_pattern = r'\b(?:\+?\d{1,4})?[\d\s\(\)-]{6,15}\b'
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    return ("✅ Email found." if emails else "❌ No email detected."), ("✅ Phone number found." if phones else "❌ No phone number detected.")

# Spell-checking
def check_spelling(text):
    spell = SpellChecker()
    words = word_tokenize(text)
    misspelled = spell.unknown(words)
    return "✅ No spelling mistakes detected." if not misspelled else f"❌ Spelling mistakes found: {', '.join(misspelled)}"

# Main function to analyze the CV
def analyze_cv(file_path, full_name):
    text = read_docx(file_path)
    
    print(check_file_type(file_path))
    print(check_last_modified_date(file_path))
    name_check, name_length_check = check_file_name(file_path, full_name)
    print(name_check)
    print(name_length_check)
    print(check_file_size(file_path))
    
    word_count, page_count = count_words_and_pages(text)
    print(f"✅ Word count: {word_count} (Ideal: 350-800)" if 350 <= word_count <= 800 else f"❌ Word count: {word_count} (Out of ideal range)")
    print(f"✅ Estimated page count: {page_count} (Ideal: 1-2 pages)" if 1 <= page_count <= 2 else f"❌ Estimated page count: {page_count} (Out of ideal range)")
    
    size_check, font_check, color_check = check_font_and_colors(file_path)
    print(size_check)
    print(font_check)
    print(color_check)
    
    print(check_sections(text))
    email_check, phone_check = check_contact_info(text)
    print(email_check)
    print(phone_check)
    
    print(check_spelling(text))
    
    return "✅ CV analysis complete with detailed feedback."


def analyze_cv_api(file_path, full_name):
    result = ""

    # Check file type and log it
    result += f"File Type: {check_file_type(file_path)}\n"
    result += f"Last Modified Date: {check_last_modified_date(file_path)}\n"

    # Check filename for the full name and length
    name_check, name_length_check = check_file_name(file_path, full_name)
    result += f"Filename Check: {name_check}\n"
    result += f"Filename Length Check: {name_length_check}\n"

    # Check file size
    result += f"File Size: {check_file_size(file_path)}\n"

    # Word count and page count checks
    text = read_docx(file_path)
    word_count, page_count = count_words_and_pages(text)
    result += f"✅ Word count: {word_count} (Ideal: 350-800)" if 350 <= word_count <= 800 else f"❌ Word count: {word_count} (Out of ideal range)\n"
    result += f"✅ Estimated page count: {page_count} (Ideal: 1-2 pages)" if 1 <= page_count <= 2 else f"❌ Estimated page count: {page_count} (Out of ideal range)\n"

    # Font and color checks
    size_check, font_check, color_check = check_font_and_colors(file_path)
    result += f"Font Size Check: {size_check}\n"
    result += f"Font Type Check: {font_check}\n"
    result += f"Font Color Check: {color_check}\n"

    # Section check (e.g., header, skills, experience)
    result += f"Sections Check: {check_sections(text)}\n"

    # Contact info check (email, phone)
    email_check, phone_check = check_contact_info(text)
    result += f"Email Check: {email_check}\n"
    result += f"Phone Check: {phone_check}\n"

    # Spelling check
    result += f"Spelling Check: {check_spelling(text)}\n"

    return result