import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_horizontal_line(doc):
    """Adds a single horizontal line to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("\u2014" * 50)  # Creates a long horizontal line
    run.bold = True

def add_section_title(doc, title_text):
    """Adds a horizontal line above and a centered, bold section title."""
    add_horizontal_line(doc)  # Move the horizontal line here, so it appears above the title

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(title_text.upper())
    run.bold = True


def add_text_section(doc, title, content, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Adds a text section with a title and content."""
    if content:
        sec_title = doc.add_paragraph(title)
        sec_title.alignment = alignment
        sec_title.runs[0].bold = True
        
        sec_content = doc.add_paragraph(content)
        sec_content.alignment = alignment

def add_skills_section(doc, skills):
    """Adds a skills section with categories in bold and items on the same line."""
    if skills:
        add_section_title(doc, "SKILLS")
        
        for skill_category in skills:
            # Add category name in bold
            category_paragraph = doc.add_paragraph()
            category_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = category_paragraph.add_run(skill_category["category"] + ": ")
            run.bold = True

            # Add the skills as a single line after the category
            skills_line = ", ".join(skill_category["items"])
            category_paragraph.add_run(skills_line)

def add_strengths_and_expertise_section(doc, strengths):
    """Adds a strengths and expertise section with bullet points."""
    if strengths:
        add_section_title(doc, "STRENGTHS AND EXPERTISE")
        for strength in strengths:
            bullet = doc.add_paragraph("• " + strength)
            bullet.paragraph_format.space_after = Pt(0)

def add_experience_section(doc, experience):
    """Adds a professional experience section."""
    if experience:
        add_section_title(doc, "PROFESSIONAL EXPERIENCE")
        
        for job in experience:
            job_title = doc.add_paragraph()
            job_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            run = job_title.add_run(job["title"] + " - " + job["company"])
            run.bold = True
            
            date_run = job_title.add_run(" " * 50 + job["dates"])
            date_run.bold = True
            
            doc.add_paragraph(job["description"])
            
            for acc in job.get("accomplishments", []):
                bullet = doc.add_paragraph("• " + acc)
                bullet.paragraph_format.space_after = Pt(0)

def create_cv_from_json(json_file, output_docx):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    
    # Name at the top
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = name.add_run(data.get("name", "Your Name").upper())
    run.bold = True
    run.font.size = Pt(14)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run(data.get("phone", "Your Phone Number") + " • ").bold = True
    contact.add_run(data.get("email", "Your Email") + " • ").bold = True
    contact.add_run(data.get("contact", "Your Location")).bold = True
    
    add_horizontal_line(doc)
    
    # Summary
    add_text_section(doc, "SOFTWARE DEVELOPER", data.get("summary", ""), WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Skills and Strengths
    add_skills_section(doc, data.get("skills", []))
    add_strengths_and_expertise_section(doc, data.get("strengths_and_expertise", []))
    
    # Experience
    add_experience_section(doc, data.get("experience", []))
    
    # Education (if needed)
    if data.get("education"):
        add_section_title(doc, "EDUCATION")
        for edu in data.get("education", []):
            doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['year']})")

    doc.save(output_docx)
    print(f"CV saved as {output_docx}")

if __name__ == "__main__":
    from ats_analysis import analyze_cv
    input_json = os.getenv("INPUT_JSON", "cv_data.json")
    output_docx = os.getenv("OUTPUT_DOCX", "output_cv.docx")
    create_cv_from_json(input_json, output_docx)
    analyze_cv("/app/output/output_cv.docx", "John Doe")
